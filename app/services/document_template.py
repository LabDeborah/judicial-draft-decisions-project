from __future__ import annotations

from copy import copy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Cm

from app.domain.types import AnalysisOutput, DocumentDecision, TnuTheme, Trf2Decision
from app.utils.fs import ensure_dir

TEMPLATE_DOCX_PATH = Path("templates/decision_template_v2.docx")
DECISION_PROCESS_LABEL = "Processo n.: "
QUOTE_LEFT_INDENT_CM = 3.1
QUOTE_RIGHT_INDENT_CM = 0.8


@dataclass(frozen=True, slots=True)
class TemplateScenario:
    key: str
    label: str
    start_paragraph: int
    end_paragraph: int
    excluded_paragraphs: tuple[int, ...] = ()


SCENARIOS = {
    "adequacao_dissonancia": TemplateScenario(
        key="adequacao_dissonancia",
        label="Adequacao por dissonancia com tema transitado",
        start_paragraph=0,
        end_paragraph=12,
        excluded_paragraphs=(1,),
    ),
    "sobrestamento_pendente_julgamento": TemplateScenario(
        key="sobrestamento_pendente_julgamento",
        label="Sobrestamento com tema ainda nao julgado",
        start_paragraph=14,
        end_paragraph=19,
        excluded_paragraphs=(14, 15),
    ),
    "sobrestamento_sem_transito": TemplateScenario(
        key="sobrestamento_sem_transito",
        label="Sobrestamento com tema julgado sem transito",
        start_paragraph=20,
        end_paragraph=29,
        excluded_paragraphs=(20,),
    ),
    "nego_seguimento_pos_sobrestamento": TemplateScenario(
        key="nego_seguimento_pos_sobrestamento",
        label="Negar seguimento apos sobrestamento",
        start_paragraph=31,
        end_paragraph=42,
        excluded_paragraphs=(32,),
    ),
    "nego_seguimento_transitado": TemplateScenario(
        key="nego_seguimento_transitado",
        label="Negar seguimento com tema ja transitado",
        start_paragraph=44,
        end_paragraph=54,
        excluded_paragraphs=(44,),
    ),
}


def template_docx_exists() -> bool:
    return TEMPLATE_DOCX_PATH.exists()


def pick_template_scenario(doc: DocumentDecision, theme: TnuTheme) -> TemplateScenario:
    if doc.action == "DETERMINAR_ADEQUACAO":
        return SCENARIOS["adequacao_dissonancia"]
    if doc.action == "SOBRESTAR":
        if not theme.dataJulgamento:
            return SCENARIOS["sobrestamento_pendente_julgamento"]
        return SCENARIOS["sobrestamento_sem_transito"]
    if doc.action == "NEGAR_SEGUIMENTO":
        # O fluxo atual nao rastreia historico processual para afirmar se houve
        # sobrestamento previo. Por padrao, usamos a variante mais generica.
        return SCENARIOS["nego_seguimento_transitado"]
    return SCENARIOS["adequacao_dissonancia"]


def render_template_docx(
    scenario: TemplateScenario,
    context: dict[str, str],
    *,
    output_path: str,
) -> str:
    document = Document(str(TEMPLATE_DOCX_PATH))
    selected_indexes = set(range(scenario.start_paragraph, scenario.end_paragraph + 1)).difference(
        scenario.excluded_paragraphs
    )
    for index in range(len(document.paragraphs) - 1, -1, -1):
        if index in selected_indexes:
            _apply_placeholder_paragraph_formatting(document.paragraphs[index])
            _replace_placeholders_in_paragraph(document.paragraphs[index], context)
            continue
        paragraph_element = document.paragraphs[index]._element
        paragraph_element.getparent().remove(paragraph_element)
    _prepend_decision_process_paragraph(document, context, scenario)
    ensure_dir(str(Path(output_path).parent))
    document.save(output_path)
    return output_path


def render_template_latex(scenario: TemplateScenario, context: dict[str, str]) -> str:
    document = Document(str(TEMPLATE_DOCX_PATH))
    selected = [
        document.paragraphs[index]
        for index in range(scenario.start_paragraph, scenario.end_paragraph + 1)
        if index not in scenario.excluded_paragraphs
    ]
    body = "\n\n".join(
        _render_latex_paragraph(paragraph, context) for paragraph in selected if _paragraph_text(paragraph)
    )
    process_header = _render_latex_decision_process_header(context)
    return (
        "\\documentclass[12pt]{article}\n"
        "\\usepackage[T1]{fontenc}\n"
        "\\usepackage[utf8]{inputenc}\n"
        "\\usepackage[brazil]{babel}\n"
        "\\usepackage[margin=1in]{geometry}\n"
        "\\usepackage{newtxtext,newtxmath}\n"
        "\\usepackage{ragged2e}\n"
        "\\usepackage[normalem]{ulem}\n"
        "\\setlength{\\parindent}{0pt}\n"
        "\\linespread{1.15}\n"
        "\\setlength{\\parskip}{0.7\\baselineskip}\n"
        "\\emergencystretch=3em\n"
        "\\sloppy\n"
        "\\begin{document}\n"
        "\\justifying\n"
        f"{process_header}{body}\n"
        "\\end{document}\n"
    )


def build_template_context(
    _doc: DocumentDecision,
    decision: Trf2Decision,
    theme: TnuTheme,
    analysis: AnalysisOutput,
) -> dict[str, str]:
    return {
        "decision_process_number": decision.numeroProcesso or "NAO_INFORMADO",
        "theme_number": theme.temaNumero or "NAO_INFORMADO",
        "theme_process_number": theme.numeroProcesso or "NAO_INFORMADO",
        "theme_affectation_date": theme.dataDecisaoAfetacao or "NAO_INFORMADO",
        "theme_judgment_date": theme.dataJulgamento or "NAO_INFORMADO",
        "theme_transit_date": theme.transitoJulgado or "NAO_INFORMADO",
        "theme_question_quote": _quote_text(theme.questaoSubmetidaJulgamento),
        "theme_thesis_quote": _quote_theme_thesis(theme),
        "theme_judgment_reference": _theme_judgment_reference(theme),
        "theme_reference_label": _theme_reference_label(theme),
        "decision_excerpt_quote": _decision_excerpt_quote(decision, analysis),
        "theme_alignment_phrase": _alignment_phrase(analysis),
    }


def list_template_placeholders() -> list[str]:
    return sorted(
        {
            "decision_process_number",
            "theme_number",
            "theme_process_number",
            "theme_affectation_date",
            "theme_judgment_date",
            "theme_transit_date",
            "theme_question_quote",
            "theme_thesis_quote",
            "theme_judgment_reference",
            "theme_reference_label",
            "decision_excerpt_quote",
            "theme_alignment_phrase",
        }
    )


def _render_latex_paragraph(paragraph, context: dict[str, str]) -> str:
    rendered_runs = [_render_latex_run(run, context) for run in paragraph.runs]
    text = "".join(rendered_runs).strip()
    if not text:
        return ""
    if paragraph.style and paragraph.style.name.startswith("Heading"):
        return f"{{\\Large\\bfseries {text}}}\\\\"
    if _is_quote_placeholder_paragraph(paragraph):
        return _render_latex_quote_paragraph(text)
    return f"\\noindent {text}"


def _render_latex_run(run, context: dict[str, str]) -> str:
    text = _replace_placeholders(run.text or "", context)
    if not text:
        return ""
    rendered = _escape_latex(text)
    if run.underline:
        rendered = f"\\uline{{{rendered}}}"
    if run.bold:
        rendered = f"\\textbf{{{rendered}}}"
    if run.italic:
        rendered = f"\\textit{{{rendered}}}"
    return rendered


def _replace_placeholders_in_paragraph(paragraph, context: dict[str, str]) -> None:
    for run in paragraph.runs:
        run.text = _replace_placeholders(run.text or "", context)


def _apply_placeholder_paragraph_formatting(paragraph) -> None:
    if not _is_quote_placeholder_paragraph(paragraph):
        return
    paragraph.paragraph_format.left_indent = Cm(QUOTE_LEFT_INDENT_CM)
    paragraph.paragraph_format.right_indent = Cm(QUOTE_RIGHT_INDENT_CM)
    paragraph.paragraph_format.first_line_indent = Cm(0)


def _is_quote_placeholder_paragraph(paragraph) -> bool:
    text = _paragraph_text(paragraph)
    return "{{theme_question_quote}}" in text or "{{theme_thesis_quote}}" in text


def _render_latex_quote_paragraph(text: str) -> str:
    return (
        "{\\leftskip=3.1cm\\rightskip=0.8cm\\parindent=0pt "
        f"{text}"
        "\\par}"
    )


def _replace_placeholders(text: str, context: dict[str, str]) -> str:
    output = text
    for key, value in context.items():
        output = output.replace(f"{{{{{key}}}}}", value)
    return output


def _prepend_decision_process_paragraph(
    document: Document,
    context: dict[str, str],
    scenario: TemplateScenario,
) -> None:
    process_number = context.get("decision_process_number", "").strip()
    if not process_number:
        return
    if any(process_number in _paragraph_text(paragraph) for paragraph in document.paragraphs):
        return
    anchor = document.paragraphs[0] if document.paragraphs else document.add_paragraph()
    reference_paragraph = _first_scenario_paragraph(document, scenario) or anchor
    reference_run = _first_text_run(reference_paragraph)
    paragraph = anchor.insert_paragraph_before()
    _copy_paragraph_format(reference_paragraph, paragraph)
    label_run = paragraph.add_run(DECISION_PROCESS_LABEL)
    label_run.bold = True
    _copy_run_format(reference_run, label_run)
    process_run = paragraph.add_run(process_number)
    process_run.underline = True
    _copy_run_format(reference_run, process_run)
    process_run.underline = True


def _first_scenario_paragraph(document: Document, scenario: TemplateScenario):
    for paragraph in document.paragraphs:
        if _paragraph_text(paragraph):
            return paragraph
    return None


def _first_text_run(paragraph):
    for run in paragraph.runs:
        if run.text:
            return run
    return paragraph.runs[0] if paragraph.runs else None


def _copy_paragraph_format(source, target) -> None:
    target.style = source.style
    target.alignment = source.alignment
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing
    target.paragraph_format.line_spacing_rule = source.paragraph_format.line_spacing_rule
    target.paragraph_format.keep_together = source.paragraph_format.keep_together
    target.paragraph_format.keep_with_next = source.paragraph_format.keep_with_next
    target.paragraph_format.page_break_before = source.paragraph_format.page_break_before
    target.paragraph_format.widow_control = source.paragraph_format.widow_control


def _copy_run_format(source, target) -> None:
    if source is None:
        return
    target.style = source.style
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    target.font.name = source.font.name
    target.font.size = source.font.size
    if source.font.color is not None and source.font.color.rgb is not None:
        target.font.color.rgb = source.font.color.rgb


def _render_latex_decision_process_header(context: dict[str, str]) -> str:
    process_number = context.get("decision_process_number", "").strip()
    if not process_number:
        return ""
    return (
        f"\\noindent \\textbf{{{_escape_latex(DECISION_PROCESS_LABEL)}}}"
        f"\\uline{{{_escape_latex(process_number)}}}\n\n"
    )


def _paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def _quote_text(text: str) -> str:
    clean = (text or "").strip().strip('"').strip("'")
    if not clean:
        return '"Nao informado."'
    return f'"{clean}"'


def _quote_theme_thesis(theme: TnuTheme) -> str:
    clean = (theme.teseFirmada or "").strip().strip('"').strip("'")
    if not clean:
        return '"Tese ainda nao firmada pela TNU."'
    suffix = f" (Tema {theme.temaNumero} TNU)." if theme.temaNumero else "."
    return f'"{clean}"{suffix}'


def _theme_judgment_reference(theme: TnuTheme) -> str:
    number = theme.temaNumero or "NAO_INFORMADO"
    date = theme.dataJulgamento or "NAO_INFORMADO"
    return f"Tema de n. {number}, em {date}"


def _theme_reference_label(theme: TnuTheme) -> str:
    number = theme.temaNumero or "NAO_INFORMADO"
    return f"Tema n. {number} da TNU"


def _decision_excerpt_quote(decision: Trf2Decision, analysis: AnalysisOutput) -> str:
    raw = (decision.assuntos or "").strip() or (analysis.justificativa or "").strip()
    if not raw:
        return "(...) Trecho relevante do acordao nao disponivel."
    if not raw.startswith("(...)"):
        raw = f"(...) {raw}"
    return raw


def _alignment_phrase(analysis: AnalysisOutput) -> str:
    if analysis.consonancia == "DISSONANCIA":
        return "o mesmo se encontra em desconformidade com o"
    return "o mesmo se encontra em conformidade com o"
def _escape_latex(text: str) -> str:
    return (
        text.replace("\\", "\\textbackslash{}")
        .replace("&", "\\&")
        .replace("%", "\\%")
        .replace("$", "\\$")
        .replace("#", "\\#")
        .replace("_", "\\_")
        .replace("{", "\\{")
        .replace("}", "\\}")
        .replace("~", "\\textasciitilde{}")
        .replace("^", "\\textasciicircum{}")
    )
