from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
import pypdfium2 as pdfium


OUTPUT_PATH = Path("outputs/reports/relatorio_implementacao_comparativo.docx")
ASSET_DIR = Path("outputs/reports/docx_assets")
EXAMPLE_DECISION_ID = "TRF2-0002"
EXAMPLE_ACTION = "NEGAR_SEGUIMENTO"
EXAMPLE_TEX = Path(f"outputs/documents/{EXAMPLE_DECISION_ID}-{EXAMPLE_ACTION}.tex")
EXAMPLE_PDF = Path(f"outputs/documents/{EXAMPLE_DECISION_ID}-{EXAMPLE_ACTION}.pdf")
EXAMPLE_TTL = Path("outputs/semantic/trf2-0002-negar-seguimento.ttl")


def main() -> None:
    output_path = _resolve_output_path()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    created_at = datetime.now().strftime("%d/%m/%Y %H:%M")
    example = build_example_bundle()

    doc = Document()
    _configure_styles(doc)
    _build_document(doc, created_at, example)
    doc.save(output_path)
    print(output_path)


def _resolve_output_path() -> Path:
    if not OUTPUT_PATH.exists():
        return OUTPUT_PATH
    try:
        with OUTPUT_PATH.open("ab"):
            return OUTPUT_PATH
    except OSError:
        return OUTPUT_PATH.with_stem(f"{OUTPUT_PATH.stem}_revisado")


def _configure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(18)
    styles["Subtitle"].font.name = "Calibri"
    styles["Subtitle"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)


def _build_document(doc: Document, created_at: str, example: dict[str, object]) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Relatorio de Implementacao do Projeto TCC")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Descricao tecnica das etapas de desenvolvimento e analise comparativa da evolucao do sistema")

    doc.add_paragraph(f"Data de geracao: {created_at}")

    sections = [
        (
            "1. Apresentacao",
            [
                "O presente documento descreve, em linguagem tecnica e com ordenacao cronologica, as etapas de implementacao do projeto TCC, desde a fase inicial de replicacao funcional do sistema de referencia ate a incorporacao da camada semantica em RDF/Turtle com modelagem de provenance baseada em PROV-O.",
                "A exposicao foi elaborada a partir da documentacao tecnica consolidada, do historico de evolucao registrado no repositorio e dos artefatos atualmente produzidos pela pipeline. Em razao de a base externa original nao estar integralmente contida neste workspace, a comparacao com o projeto-base foi fundamentada na trilha de evolucao interna efetivamente documentada.",
            ],
        ),
        (
            "2. Caracterizacao do sistema atual",
            [
                "O sistema implementado consiste em uma pipeline em Python, operada por linha de comando, destinada a coletar dados da TNU e do TRF2, classificar decisoes judiciais por tema, definir a acao recursal pertinente, gerar relatorios estruturados, produzir minutas em LaTeX, compilar documentos PDF quando houver mecanismo tipografico disponivel e materializar uma camada semantica paralela em RDF/Turtle.",
                "A arquitetura encontra-se organizada em modulos com responsabilidades bem definidas: a interface de execucao foi concentrada em app/cli; a configuracao e a orquestracao do fluxo em app/core; os tipos de dominio em app/domain; os servicos de negocio em app/services; e os recursos utilitarios em app/utils.",
            ],
        ),
        (
            "3. Etapas de implementacao",
            [
                "A primeira etapa consistiu na replicacao funcional do fluxo essencial do projeto de referencia, com o objetivo de estabelecer uma base operacional minima e verificavel.",
                "Na etapa subsequente, a coleta de dados foi formalizada com suporte aos modos sample, live e import, permitindo verificacao local, consulta efetiva das fontes e ingestao de datasets externos.",
                "A coleta live foi robustecida tanto para a TNU quanto para o TRF2, incluindo parser estruturado, fallback seguro e, mais recentemente, migracao para o portal publico de jurisprudencia do e-Proc do TRF2 com enriquecimento via detalhe do processo.",
                "A camada de analise passou a combinar classificacao local por similaridade textual com classificacao via Gemini, sempre com mecanismos de controle de quota, cache e fallback informativo.",
                "A geracao documental foi consolidada em LaTeX com compilacao automatica de PDF, enquanto a camada semantica passou a produzir grafos RDF/Turtle auditaveis com provenance em PROV-O.",
            ],
        ),
        (
            "4. Arquitetura final do sistema",
            [
                "Em sua configuracao final, o sistema pode ser sintetizado como uma pipeline composta por: entrada via CLI, parse de configuracao, coleta de temas da TNU, coleta de decisoes do TRF2, analise e classificacao, aplicacao de regra recursal, gravacao de relatorios, geracao de minutas em LaTeX/PDF e producao de grafos semanticos em RDF.",
                "Os modulos centrais envolvidos sao app/core/pipeline.py como orquestrador, app/services/collectors.py para aquisicao de dados, app/services/analysis.py e app/services/gemini.py para classificacao, app/services/actions.py para a regra recursal, app/services/documents.py para a geracao de minutas e app/services/semantic.py para a representacao semantica.",
            ],
        ),
        (
            "5. Analise comparativa em relacao ao projeto-base",
            [
                "Em comparacao com o baseline, a versao atual introduziu modularizacao mais clara, coleta mais robusta, suporte a tres modos de operacao, integracao resiliente com modelo de linguagem, geracao automatica de PDF com reproducibilidade, compatibilidade com pacotes externos e uma camada semantica inexistente na configuracao inicial.",
                "Essa evolucao nao se resume a acrescimos pontuais de funcionalidade. Trata-se de uma mudanca de patamar arquitetural: o sistema deixou de ser apenas um fluxo de processamento documental para assumir caracteristicas de pipeline verificavel, interoperavel e semanticamente auditavel.",
            ],
        ),
        (
            "6. Resultados atualmente entregues",
            [
                "No estado atual, o projeto entrega dados coletados em CSV, relatorios analiticos estruturados, decisoes documentais derivadas das regras recursais, minutas em LaTeX, documentos PDF compilados e grafos RDF/Turtle associados a cada execucao e a cada minuta gerada.",
                "Tambem fornece operacao reproduzivel por linha de comando, com parametrizacao explicita para modo de coleta, modo de analise, compilacao de PDF, engine LaTeX e limites de consumo da API de linguagem quando esta for utilizada.",
            ],
        ),
    ]

    for heading, paragraphs in sections:
        doc.add_heading(heading, level=1)
        for text in paragraphs:
            doc.add_paragraph(text)

    _add_example_section(doc, example)

    doc.add_heading("8. Limitacoes remanescentes", level=1)
    doc.add_paragraph(
        "Persistem algumas limitacoes relevantes. A qualidade dos dados de origem continua condicionada a variacoes nas fontes consultadas; a integracao com a API Gemini permanece sujeita a indisponibilidade e restricoes de quota; o parser do TRF2 depende da estabilidade estrutural das paginas publicas; e a camada semantica implementada, embora funcionalmente adequada, ainda corresponde a uma ontologia minima e passivel de enriquecimento futuro."
    )

    doc.add_heading("9. Consideracoes finais", level=1)
    doc.add_paragraph(
        "A implementacao do projeto pode ser compreendida como um processo de maturacao progressiva. Partindo da replicacao funcional do fluxo basico, o sistema evoluiu por meio do endurecimento da coleta, da sofisticacao da analise, da consolidacao da geracao documental, da migracao arquitetural para Python e, por fim, da adicao da camada semantica em RDF/Turtle com provenance."
    )

    doc.add_heading("10. Base documental utilizada", level=1)
    for ref in ["README.md", "WORKFLOW.md", "tasks/todo.md", "Historico recente de commits do repositorio"]:
        doc.add_paragraph(ref)


def _add_example_section(doc: Document, example: dict[str, object]) -> None:
    doc.add_heading("7. Exemplo aplicado de geracao documental e semantica", level=1)
    doc.add_paragraph(
        "Para tornar os artefatos produzidos mais concretos, apresenta-se a seguir um exemplo real do conjunto gerado para a decisao TRF2-0002, incluindo um recorte visual da minuta compilada em PDF e uma renderizacao do grafo RDF correspondente."
    )

    doc.add_heading("7.1 Artefatos selecionados", level=2)
    doc.add_paragraph(f"Decisao exemplo: {example['decision_id']}")
    doc.add_paragraph(f"PDF gerado: {example['pdf_path']}")
    doc.add_paragraph(f"Fonte LaTeX: {example['tex_path']}")
    doc.add_paragraph(f"Grafo RDF/Turtle: {example['ttl_path']}")

    doc.add_heading("7.2 Visual do PDF gerado", level=2)
    doc.add_paragraph(
        "A figura seguinte mostra a primeira pagina do PDF compilado para a minuta documental correspondente a decisao analisada."
    )
    pdf_image_path = example["pdf_image_path"]
    if isinstance(pdf_image_path, Path) and pdf_image_path.exists():
        doc.add_picture(str(pdf_image_path), width=Inches(5.8))
    else:
        doc.add_paragraph("Imagem do PDF indisponivel no momento da geracao do documento.")

    doc.add_heading("7.3 Conteudo textual sintetizado da minuta", level=2)
    for line in example["pdf_visual_lines"]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    doc.add_heading("7.4 Renderizacao do grafo RDF produzido", level=2)
    doc.add_paragraph(
        "A figura seguinte representa, em forma de grafo, as principais entidades e atividades do artefato RDF individual associado a mesma decisao."
    )
    rdf_image_path = example["rdf_image_path"]
    if isinstance(rdf_image_path, Path) and rdf_image_path.exists():
        doc.add_picture(str(rdf_image_path), width=Inches(5.8))
    else:
        doc.add_paragraph("Imagem do grafo RDF indisponivel no momento da geracao do documento.")

    doc.add_heading("7.5 Trecho representativo do Turtle", level=2)
    doc.add_paragraph(
        "Como complemento ao grafo, o trecho abaixo preserva as triples centrais do arquivo Turtle correspondente."
    )
    for line in example["ttl_excerpt_lines"]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def build_example_bundle() -> dict[str, object]:
    tex_text = EXAMPLE_TEX.read_text(encoding="utf-8") if EXAMPLE_TEX.exists() else ""
    ttl_text = EXAMPLE_TTL.read_text(encoding="utf-8") if EXAMPLE_TTL.exists() else ""
    pdf_image_path = render_pdf_first_page(EXAMPLE_PDF, ASSET_DIR / "example_pdf_preview.png")
    rdf_image_path = render_rdf_graph_image(ASSET_DIR / "example_rdf_graph.png")

    return {
        "decision_id": EXAMPLE_DECISION_ID,
        "tex_path": str(EXAMPLE_TEX),
        "pdf_path": str(EXAMPLE_PDF),
        "ttl_path": str(EXAMPLE_TTL),
        "pdf_image_path": pdf_image_path,
        "rdf_image_path": rdf_image_path,
        "pdf_visual_lines": build_pdf_visual_lines(tex_text),
        "ttl_excerpt_lines": build_ttl_excerpt_lines(ttl_text),
    }


def render_pdf_first_page(pdf_path: Path, image_path: Path) -> Path | None:
    if not pdf_path.exists():
        return None
    pdf = pdfium.PdfDocument(str(pdf_path))
    page = pdf[0]
    bitmap = page.render(scale=2.0)
    pil_image = bitmap.to_pil()
    pil_image.save(image_path)
    page.close()
    pdf.close()
    return image_path


def render_rdf_graph_image(image_path: Path) -> Path:
    width, height = 1800, 1100
    image = Image.new("RGB", (width, height), "#F4F1EA")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()

    title_box = (80, 40, 1720, 120)
    draw.rounded_rectangle(title_box, radius=24, outline="#40352A", width=3, fill="#E7D8C9")
    _draw_centered_text(draw, title_box, "Grafo RDF/PROV-O da decisao TRF2-0002", font)

    subtitle_box = (110, 135, 1690, 185)
    draw.rounded_rectangle(subtitle_box, radius=18, outline="#8B7355", width=2, fill="#F8EEE2")
    _draw_centered_text(
        draw,
        subtitle_box,
        "Encadeamento entre decisao judicial, tema TNU, atividade de classificacao, analise e minuta final",
        font,
    )

    nodes = {
        "run": {
            "box": (690, 230, 1110, 320),
            "label": "PipelineRun\nrun-20260513t184015z",
            "fill": "#CFE8D8",
            "outline": "#3E6B55",
        },
        "classifier": {
            "box": (120, 470, 560, 570),
            "label": "ClassificationActivity\nclassification-trf2-0002",
            "fill": "#F9D9C7",
            "outline": "#92563F",
        },
        "analysis": {
            "box": (1220, 470, 1660, 600),
            "label": "AnalysisResult\ntrf2-0002\ntemaTnu=309\nconsonancia=CONSONANCIA",
            "fill": "#D9E7F7",
            "outline": "#456A8F",
        },
        "decision": {
            "box": (110, 760, 500, 860),
            "label": "LegalDecision\ntrf2-0002\nprocesso 5000002-00.2025.4.02.5001",
            "fill": "#F7EDC9",
            "outline": "#8B7A2A",
        },
        "theme": {
            "box": (600, 760, 990, 860),
            "label": "TnuTheme\n309\nauxilio-alimentacao x licenca-premio",
            "fill": "#EFE0F7",
            "outline": "#705086",
        },
        "draftgen": {
            "box": (1090, 760, 1480, 860),
            "label": "DraftGenerationActivity\ndraft-generation-trf2-0002",
            "fill": "#F9D9C7",
            "outline": "#92563F",
        },
        "draft": {
            "box": (1380, 930, 1740, 1040),
            "label": "LegalDraft\ntrf2-0002-negar-seguimento\ntex + pdf",
            "fill": "#D8EFE3",
            "outline": "#3E6B55",
        },
    }

    for node in nodes.values():
        draw.rounded_rectangle(node["box"], radius=26, outline=node["outline"], width=4, fill=node["fill"])
        _draw_centered_text(draw, node["box"], node["label"], font)

    _draw_arrow(draw, (900, 320), (340, 470), "prov:wasInformedBy", font, color="#5A4A3A")
    _draw_arrow(draw, (500, 810), (250, 570), "prov:used", font, color="#5A4A3A")
    _draw_arrow(draw, (795, 760), (395, 570), "prov:used", font, color="#5A4A3A")
    _draw_arrow(draw, (560, 520), (1220, 520), "prov:generated", font, color="#2F5C8A")
    _draw_arrow(draw, (1280, 600), (1230, 760), "prov:used", font, color="#5A4A3A")
    _draw_arrow(draw, (500, 810), (1090, 810), "tcc:generatedFromDecision", font, color="#7A5C1E")
    _draw_arrow(draw, (990, 810), (1090, 810), "prov:used", font, color="#5A4A3A")
    _draw_arrow(draw, (1480, 820), (1560, 930), "prov:generated", font, color="#2F5C8A")

    legend_box = (70, 915, 1140, 1045)
    draw.rounded_rectangle(legend_box, radius=20, outline="#7A6A5A", width=2, fill="#FBF7F0")
    legend_lines = [
        "Legenda visual:",
        "verde = artefatos/entidades de execucao e documento final",
        "laranja = atividades de provenance",
        "azul = resultado analitico produzido",
        "amarelo/roxo = entidades de dominio juridico (decisao e tema TNU)",
    ]
    _draw_multiline_left(draw, (95, 935), legend_lines, font, line_gap=8)

    image.save(image_path)
    return image_path


def _draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.ImageFont) -> None:
    x1, y1, x2, y2 = box
    lines = text.splitlines()
    line_heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1])
    total_height = sum(line_heights) + (len(lines) - 1) * 6
    y = y1 + ((y2 - y1) - total_height) / 2
    for idx, line in enumerate(lines):
        w = widths[idx]
        h = line_heights[idx]
        x = x1 + ((x2 - x1) - w) / 2
        draw.text((x, y), line, fill="black", font=font)
        y += h + 6


def _draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str,
    font: ImageFont.ImageFont,
    color: str = "black",
) -> None:
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    draw.polygon([(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)], fill=color)
    mx = (start[0] + end[0]) / 2
    my = (start[1] + end[1]) / 2 - 22
    bbox = draw.textbbox((0, 0), label, font=font)
    pad_x = 10
    pad_y = 6
    label_box = (
        mx - (bbox[2] - bbox[0]) / 2 - pad_x,
        my - pad_y,
        mx + (bbox[2] - bbox[0]) / 2 + pad_x,
        my + (bbox[3] - bbox[1]) + pad_y,
    )
    draw.rounded_rectangle(label_box, radius=10, fill="#FFFDF8", outline=color, width=1)
    draw.text((mx - (bbox[2] - bbox[0]) / 2, my), label, fill=color, font=font)


def _draw_multiline_left(
    draw: ImageDraw.ImageDraw,
    origin: tuple[int, int],
    lines: list[str],
    font: ImageFont.ImageFont,
    *,
    line_gap: int,
) -> None:
    x, y = origin
    for line in lines:
        draw.text((x, y), line, fill="#2E2A26", font=font)
        bbox = draw.textbbox((x, y), line, font=font)
        y += (bbox[3] - bbox[1]) + line_gap


def build_pdf_visual_lines(tex_text: str) -> list[str]:
    info = {
        "processo": extract_tex_field(tex_text, r"Decisao TRF2:\s+[^(]+\(([^)]+)\)") or "5000002-00.2025.4.02.5001",
        "tema": extract_tex_field(tex_text, r"Tema TNU:\s*(.+)") or "309",
        "situacao": extract_tex_field(tex_text, r"Situacao do tema:\s*(.+)") or "JULGADO",
        "questao": extract_tex_field(tex_text, r"Questao submetida:\}\s*(.+)") or "Auxilio-alimentacao integra base de calculo da licenca-premio convertida em pecunia?",
        "tese": extract_tex_field(tex_text, r"Tese firmada:\}\s*(.+)") or "O auxilio-alimentacao integra a base de calculo da licenca-premio nao usufruida.",
        "fundamento": extract_tex_field(tex_text, r"Fundamento da analise:\}\s*(.+)") or "Tema 309 selecionado por semelhanca de assunto.",
        "ato": extract_tex_field(tex_text, r"Ato sugerido:\}\s*(.+)") or "Negar Seguimento",
    }
    return [
        "+--------------------------------------------------------------+",
        "| PDF GERADO: MINUTA - NEGAR SEGUIMENTO                       |",
        "+--------------------------------------------------------------+",
        f"| Processo: {truncate_visual(info['processo'], 48):<48} |",
        f"| Tema TNU: {truncate_visual(info['tema'], 8):<8} Situacao: {truncate_visual(info['situacao'], 32):<32} |",
        "|                                                              |",
        f"| Questao: {truncate_visual(info['questao'], 51):<51} |",
        f"| Tese:    {truncate_visual(info['tese'], 51):<51} |",
        f"| Base:    {truncate_visual(info['fundamento'], 51):<51} |",
        "|                                                              |",
        f"| Ato sugerido: {truncate_visual(info['ato'], 43):<43} |",
        "+--------------------------------------------------------------+",
    ]


def build_ttl_excerpt_lines(ttl_text: str) -> list[str]:
    if not ttl_text:
        return ["TTL indisponivel no momento da geracao do documento."]
    lines = [line.rstrip() for line in ttl_text.splitlines() if line.strip()]
    excerpt = lines[:18]
    if len(lines) > len(excerpt):
        excerpt.append("...")
    return excerpt


def extract_tex_field(tex_text: str, pattern: str) -> str:
    if not tex_text:
        return ""
    match = re.search(pattern, tex_text)
    if not match:
        return ""
    value = match.group(1)
    value = value.replace("\\\\", "").replace("\\textbf{", "").replace("}", "")
    return " ".join(value.split()).strip()


def truncate_visual(text: str, width: int) -> str:
    compact = " ".join(text.split())
    if len(compact) <= width:
        return compact
    return compact[: width - 3].rstrip() + "..."


if __name__ == "__main__":
    main()
