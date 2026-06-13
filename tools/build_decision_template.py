from __future__ import annotations

from pathlib import Path
import os
import sys

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from app.utils.fs import ensure_dir

SOURCE_DOCX = Path(r"C:\Users\ejmar\Downloads\decisoes para Hygor.docx")
OUTPUT_DOCX = Path("templates/decision_template_v2.docx")
README_PATH = Path("templates/README.md")

HIGHLIGHT_GROUPS = [
    ((2, (2,)), "{{theme_number}}"),
    ((2, (4,)), "{{theme_process_number}}"),
    ((2, (6,)), " {{theme_affectation_date}}"),
    ((2, (8,)), "{{theme_transit_date}}"),
    ((3, (0,)), "{{theme_question_quote}}"),
    ((5, (0, 1, 2)), "{{theme_thesis_quote}}"),
    ((6, (1,)), "{{theme_alignment_phrase}}"),
    ((7, (1,)), "{{decision_excerpt_quote}}"),
    ((16, (1,)), "{{theme_number}} "),
    ((16, (3,)), "{{theme_process_number}}"),
    ((16, (5,)), " {{theme_affectation_date}}"),
    ((17, (0,)), "{{theme_question_quote}}"),
    ((18, (5,)), " {{theme_process_number}}"),
    ((18, (8,)), "{{theme_number}} "),
    ((21, (1,)), "{{theme_process_number}}"),
    ((21, (3,)), "{{theme_affectation_date}}"),
    ((21, (5, 6, 7)), "{{theme_question_quote}}"),
    ((23, (0,)), "{{theme_thesis_quote}}"),
    ((24, (1,)), "{{theme_judgment_reference}}"),
    ((28, (10,)), "{{theme_process_number}}"),
    ((28, (12,)), "{{theme_reference_label}}"),
    ((33, (1,)), "{{theme_number}} "),
    ((33, (3,)), "{{theme_process_number}}"),
    ((33, (5,)), "{{theme_judgment_date}}"),
    ((33, (7,)), "{{theme_transit_date}}"),
    ((34, (0,)), "{{theme_question_quote}}"),
    ((35, (1,)), "{{theme_number}}"),
    ((36, (0,)), "{{theme_thesis_quote}}"),
    ((38, (0,)), "{{decision_excerpt_quote}}"),
    ((45, (2,)), "{{theme_number}}"),
    ((45, (4,)), "{{theme_process_number}}"),
    ((45, (6,)), " {{theme_affectation_date}}"),
    ((45, (8,)), "{{theme_transit_date}}"),
    ((46, (0,)), "{{theme_question_quote}}"),
    ((48, (0, 1, 2)), "{{theme_thesis_quote}}"),
    ((49, (1,)), "{{theme_alignment_phrase}}"),
    ((50, (1,)), "{{decision_excerpt_quote}}"),
]


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"Documento de referencia nao encontrado: {SOURCE_DOCX}")
    ensure_dir(str(OUTPUT_DOCX.parent))
    document = Document(str(SOURCE_DOCX))
    _replace_highlights_with_placeholders(document)
    _neutralize_reference_notes(document)
    temp_output = OUTPUT_DOCX.with_suffix(".tmp.docx")
    document.save(str(temp_output))
    os.replace(temp_output, OUTPUT_DOCX)
    README_PATH.write_text(_build_readme(), encoding="utf-8")
    print(f"Template DOCX gerado em {OUTPUT_DOCX}")
    print(f"Guia de placeholders gerado em {README_PATH}")


def _neutralize_reference_notes(document: Document) -> None:
    document.paragraphs[1].text = ""
    document.paragraphs[15].text = ""
    document.paragraphs[20].text = ""
    document.paragraphs[32].text = ""
    document.paragraphs[44].text = ""
    document.paragraphs[7].runs[0].text = "A proposito, consta do acordao hostilizado: "
    document.paragraphs[21].text = (
        "Verifica-se que a questao objeto da controversia juridica suscitada pelo recurso manejado pela "
        "parte recorrente encontra-se afetada ao Tema de n. {{theme_number}} da jurisprudencia da TNU "
        "- PU no processo n. {{theme_process_number}}, afetado em {{theme_affectation_date}}, voltado "
        "a elucidar a seguinte questao: {{theme_question_quote}}"
    )
    document.paragraphs[37].text = "O acordao recorrido assentou:"
    document.paragraphs[50].runs[0].text = "A proposito, consta do acordao hostilizado: "


def _replace_highlights_with_placeholders(document: Document) -> None:
    for (paragraph_index, run_indexes), placeholder in HIGHLIGHT_GROUPS:
        paragraph = document.paragraphs[paragraph_index]
        first_run = paragraph.runs[run_indexes[0]]
        first_run.text = placeholder
        first_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        for extra_run_index in run_indexes[1:]:
            extra_run = paragraph.runs[extra_run_index]
            extra_run.text = ""
            extra_run.font.highlight_color = None
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if "{{" not in run.text and run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                run.font.highlight_color = None


def _build_readme() -> str:
    return """# Template documental

Arquivos:
- `decision_template_v2.docx`: template principal derivado do documento de referencia.

Cenarios estruturais:
- `adequacao_dissonancia`: paragrafos 1-13
- `sobrestamento_pendente_julgamento`: paragrafos 15-20
- `sobrestamento_sem_transito`: paragrafos 21-30
- `nego_seguimento_pos_sobrestamento`: paragrafos 32-43
- `nego_seguimento_transitado`: paragrafos 45-55

Placeholders principais:
- `{{theme_number}}`
- `{{theme_process_number}}`
- `{{theme_affectation_date}}`
- `{{theme_judgment_date}}`
- `{{theme_transit_date}}`
- `{{theme_question_quote}}`
- `{{theme_thesis_quote}}`
- `{{theme_judgment_reference}}`
- `{{theme_reference_label}}`
- `{{decision_excerpt_quote}}`
- `{{theme_alignment_phrase}}`
"""


if __name__ == "__main__":
    main()
